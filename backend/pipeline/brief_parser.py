"""Brief parser — extracts structured fields from a .docx upload.

Uses python-docx to pull text from paragraphs and tables. Text boxes and
embedded OLE objects cannot be reliably extracted, so the parser flags
their presence for Sofie to warn the user.
"""

from pathlib import Path

from docx import Document

from backend.schemas import BriefParseResult

# Sections map to the brief template defined in PRD.md section 5.
# Keys are normalised lowercase — matched against heading text.
_SECTION_HEADINGS: dict[str, str] = {
    "brand information": "brand_info",
    "job details": "job_details",
    "campaign context": "campaign",
    "asset links": "assets",
    "layout references": "layout_refs",
    "copy": "copy",
    "restrictions": "restrictions",
}

# Field labels that appear in the template, mapped to output keys.
_FIELD_MAP: dict[str, str] = {
    "brand name": "brand_name",
    "industry": "industry",
    "brand guidelines link": "brand_guidelines_link",
    "job title": "job_title",
    "platform": "platform",
    "platforms": "platform",
    "output sizes": "output_sizes",
    "output size": "output_sizes",
    "campaign objective": "campaign_objective",
    "key message": "key_message",
    "call to action": "cta_text",
    "call to action text": "cta_text",
    "cta": "cta_text",
    "logo link": "logo_link",
    "logo": "logo_link",
    "brand font link": "brand_font_link",
    "font link": "brand_font_link",
    "hero image link": "hero_image_links",
    "hero image links": "hero_image_links",
    "hero image": "hero_image_links",
    "design elements link": "design_elements_links",
    "design elements links": "design_elements_links",
    "design elements": "design_elements_links",
    "brand colour palette": "brand_colours",
    "brand colors": "brand_colours",
    "brand colours": "brand_colours",
    "colour palette": "brand_colours",
    "own past ad link": "own_past_ad_links",
    "own past ad links": "own_past_ad_links",
    "external layout reference link": "external_ref_links",
    "external layout reference links": "external_ref_links",
    "mood reference link": "mood_ref_links",
    "mood reference links": "mood_ref_links",
    "headline text": "headline_text",
    "headline": "headline_text",
    "sub-copy": "sub_copy",
    "sub copy": "sub_copy",
    "subcopy": "sub_copy",
    "mandatory inclusions": "mandatory_inclusions",
    "what not to do": "restrictions_dont",
    "colours to avoid": "colours_to_avoid",
    "colors to avoid": "colours_to_avoid",
    "elements to avoid": "elements_to_avoid",
}

# Fields that can hold multiple values (comma or newline separated).
_LIST_FIELDS: set[str] = {
    "output_sizes",
    "hero_image_links",
    "design_elements_links",
    "own_past_ad_links",
    "external_ref_links",
    "mood_ref_links",
}


async def parse_brief(docx_path: Path) -> BriefParseResult:
    """Extract structured fields from a .docx brief.

    Paragraphs are scanned for known field labels (case-insensitive).
    Tables are scanned row-by-row for label:value pairs.
    The parser intentionally over-matches so downstream validation
    (Priya) can decide what's missing vs what's truly absent.
    """
    doc = Document(str(docx_path))
    fields: dict[str, str | list[str] | None] = {}
    warnings: list[str] = []
    has_text_boxes = False

    # Check for text boxes / embedded objects that python-docx cannot read.
    # These live in the XML body as w:txbxContent or mc:AlternateContent.
    xml_body = doc.element.body.xml
    if "w:txbxContent" in xml_body or "mc:AlternateContent" in xml_body:
        has_text_boxes = True
        warnings.append(
            "Brief contains text boxes or embedded objects that could not be "
            "fully extracted. Please verify the extracted fields are complete."
        )

    # Extract from paragraphs — look for "Label: Value" patterns
    _extract_from_paragraphs(doc, fields)

    # Extract from tables — each row treated as label:value
    _extract_from_tables(doc, fields, warnings)

    return BriefParseResult(
        fields=fields,
        warnings=warnings,
        has_text_boxes=has_text_boxes,
    )


def _extract_from_paragraphs(
    doc: Document, fields: dict[str, str | list[str] | None]
) -> None:
    """Scan paragraphs for 'Label: Value' patterns.

    Heading paragraphs are skipped — they denote sections, not fields.
    Colon is the primary delimiter; em-dash and tab are fallbacks.
    """
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Skip headings — they're section markers, not field values
        if para.style and para.style.name and "Heading" in para.style.name:
            continue

        _try_extract_field(text, fields)


def _extract_from_tables(
    doc: Document,
    fields: dict[str, str | list[str] | None],
    warnings: list[str],
) -> None:
    """Scan table rows for label:value pairs.

    Two-column tables are the most common brief format — column 0 is
    the label, column 1 is the value. Single-column tables fall back
    to the colon-delimited parser.
    """
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 2 and cells[0] and cells[1]:
                _try_extract_field(f"{cells[0]}: {cells[1]}", fields)
            elif len(cells) == 1 and cells[0]:
                _try_extract_field(cells[0], fields)


def _try_extract_field(
    text: str, fields: dict[str, str | list[str] | None]
) -> None:
    """Attempt to parse a single text line into a known field.

    Matches the text prefix against _FIELD_MAP keys. If the field is
    a list type, values are split on commas or newlines.
    """
    # Try colon separator first, then em-dash
    for sep in (":", "—", "–", "\t"):
        if sep in text:
            parts = text.split(sep, 1)
            label = parts[0].strip().lower().rstrip("*")
            value = parts[1].strip() if len(parts) > 1 else ""

            if label in _FIELD_MAP and value:
                field_key = _FIELD_MAP[label]
                if field_key in _LIST_FIELDS:
                    # Split comma-separated or newline-separated values
                    items = [
                        v.strip()
                        for v in value.replace("\n", ",").split(",")
                        if v.strip()
                    ]
                    existing = fields.get(field_key)
                    if isinstance(existing, list):
                        existing.extend(items)
                    else:
                        fields[field_key] = items
                else:
                    fields[field_key] = value
                return
